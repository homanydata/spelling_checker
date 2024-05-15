import os
from docx import Document
from helper import DATASET_PATH, ERRORS


"""
This file is responsible for interaction with text and docx files
- read dataset text file
- read existing file (text or word doc)
- save to a text file
- save to word doc file
- apply suggestions to a file

"""


def get_extension(file_path: str) -> str:
    return file_path.split('.')[-1]


def get_parent_path(file_path: str) -> str:
    return os.path.dirname(file_path)


def get_file_name(file_path: str) -> str:
    return os.path.basename(file_path).split('.')[0]


def is_document(file_path: str) -> bool:
    extension = get_extension(file_path)
    return extension in ('docx', 'doc')


def read_dataset() -> list[str]:
    contents = read_file(DATASET_PATH)
    words = contents.lower().replace('\n',' ').split(' ')
    return words


def read_file(file_path: str) -> str:
    try:
        if is_document(file_path):
            doc = Document(file_path)
            contents = '\n\n'.join((paragraph.text for paragraph in doc.paragraphs))
        else:
            with open(file_path, 'r', encoding='utf-8') as file:
                contents = file.read()
        return contents
    except:
        raise RuntimeError(ERRORS.ERROR_READING_FILE(file_path))


def save_to_file(file_path: str, text: str) -> bool:
    try:
        if is_document(file_path):
            save_to_docx_file(file_path, text)
        else:
            save_to_text_file(file_path, text)
    except:
        raise RuntimeError(ERRORS.ERROR_SAVING_FILE(file_path))



# not to be used outside of this module, just helpers to simpify stuff
def save_to_text_file(file_path:str, text: str) -> bool:
    with open(file_path, 'w', encoding='utf-8') as file:
        file.write(text)


def save_to_docx_file(file_path:str, text: str) -> bool:
    doc = Document()
    doc.add_paragraph(text)
    doc.save(file_path)
